import os
import tkinter as tk
from tkinter import filedialog, scrolledtext
import fitz  # PyMuPDF
import qianfan
from docx import Document  # Add this import

# Paste the assignment submission here （KEY）
os.environ["QIANFAN_ACCESS_KEY"] = "b6620cfa76414f9ca667330bc8966f4e"
os.environ["QIANFAN_SECRET_KEY"] = "6eb4d707b54044858d73cf73eb69ccdb"

def extract_first_page_text(pdf_path):
    doc = fitz.open(pdf_path)
    page = doc.load_page(0)
    text = page.get_text()
    return text

def process_pdfs_in_folder(folder_path):
    result_text = ""
    pdf_texts = {}
    for filename in os.listdir(folder_path):
        if filename.lower().endswith('.pdf'):
            pdf_path = os.path.join(folder_path, filename)
            result_text += f"Processing: {filename}\n"
            try:
                text = extract_first_page_text(pdf_path)
                pdf_texts[filename] = text
                result_text += f"Extracted text from the first page of {filename}\n"
                result_text += "-" * 50 + "\n"
            except Exception as e:
                result_text += f"Error processing {filename}: {str(e)}\n"
    return result_text, pdf_texts

def select_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        folder_entry.delete(0, tk.END)
        folder_entry.insert(0, folder_path)

def process_pdfs():
    folder_path = folder_entry.get()
    if os.path.exists(folder_path):
        global pdf_texts
        result, pdf_texts = process_pdfs_in_folder(folder_path)
        result_text.delete('1.0', tk.END)
        result_text.insert(tk.END, result)
    else:
        result_text.delete('1.0', tk.END)
        result_text.insert(tk.END, "The specified path does not exist. Please check and try again.")

def send_to_baidu():
    if not pdf_texts:
        result_text.insert(tk.END, "\nPlease process the PDF files first.\n")
        return None

    result_text.insert(tk.END, "\nGenerating APA citations using Baidu AI...\n")
    results = []

    chat_comp = qianfan.ChatCompletion()
    document = Document()  # Create a new Document

    with open("apa_captions.txt", "w", encoding="utf-8") as file:
        for filename, text in pdf_texts.items():
            try:
                prompt = f"Generate an APA citation based on the following text from the first page of the PDF:\n{text}\n\nAPA citation:"
                response = chat_comp.do(model="ERNIE-4.0-8K-Latest", messages=[{
                    "role": "user",
                    "content": prompt
                }])
                response_json = response["body"]
                print(response_json['result'])

                if isinstance(response_json, dict) and 'result' in response_json and response_json['result']:
                    result = response_json['result']
                    result_text.insert(tk.END, f"\nBaidu Wenxin Yiyan APA citation for '{filename}':\n")
                    result_text.insert(tk.END, str(result) + "\n")

                    file.write(f"APA citation for {filename}:\n")
                    file.write(str(result) + "\n\n")
                    results.append(result)

                    # Write to Word document
                    document.add_heading(f"APA citation for {filename}:", level=1)
                    document.add_paragraph(str(result))
                else:
                    raise ValueError("API response does not contain 'result' or it is empty.")
            except Exception as e:
                result_text.insert(tk.END, f"\nError processing '{filename}': {str(e)}\n")
                file.write(f"Error processing '{filename}': {str(e)}\n\n")
                results.append(str(e))

    # Save the Word document
    document.save("apa_captions.docx")

    return results

# Create main window
root = tk.Tk()
root.title("APA citation generator")
root.geometry("800x600")

# Create and place folder selection section
folder_frame = tk.Frame(root)
folder_frame.pack(pady=10)

folder_label = tk.Label(folder_frame, text="PDF Folder Path:")
folder_label.pack(side=tk.LEFT)

folder_entry = tk.Entry(folder_frame, width=50)
folder_entry.pack(side=tk.LEFT, padx=5)

folder_button = tk.Button(folder_frame, text="Select Folder", command=select_folder)
folder_button.pack(side=tk.LEFT)

# Create and place buttons
button_frame = tk.Frame(root)
button_frame.pack(pady=10)

process_button = tk.Button(button_frame, text="Process PDF Files", command=process_pdfs)
process_button.pack(side=tk.LEFT, padx=5)

baidu_button = tk.Button(button_frame, text="Generate APA Citations", command=send_to_baidu)
baidu_button.pack(side=tk.LEFT, padx=5)

# Create and place result display area
result_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=80, height=30)
result_text.pack(padx=10, pady=10)

# Global variable to store PDF texts
pdf_texts = {}

# Run main loop
root.mainloop()